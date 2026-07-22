from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.style = doc.styles['Normal']

def add(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    return p

def read_src(filename):
    path = os.path.join(base, filename)
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()

# Maven standard directory: src/main/java
base = r'C:\Users\DeathAngel\Documents\Java\Creational Design Pattern\src\main\java'

# ===== TITLE =====
for _ in range(6):
    doc.add_paragraph()
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('Design Patterns\nPractical Report')
r.font.size = Pt(24); r.bold = True; r.font.name = 'Times New Roman'
doc.add_paragraph()
sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run('Singleton  |  Factory Method  |  Builder  |  Adapter  |  Facade  |  Proxy  |  Decorator\n\n'
               'Student Admission Module  |  Student Services Module')
r.font.size = Pt(14); r.font.name = 'Times New Roman'
doc.add_page_break()

# ===== TOC =====
doc.add_heading('Table of Contents', level=1)
for item in [
    'Practical 1: Singleton Pattern',
    'Practical 2: Factory Method Pattern',
    'Practical 3: Builder Pattern',
    'Practical 4: Adapter Pattern',
    'Practical 5: Facade Pattern',
    'Practical 6: Proxy Pattern',
    'Practical 7: Decorator Pattern',
    'Assignment 1: Student Admission Mini Module',
    'Assignment 2: Student Services Mini Module',
]:
    add(doc, item)
doc.add_page_break()

# ===== PRACTICAL 1: SINGLETON =====
doc.add_heading('Practical 1: Singleton Pattern', level=1)
doc.add_heading('Objective', level=2)
for o in [
    'Identify problems caused by creating multiple objects unnecessarily.',
    'Understand the need for the Singleton Design Pattern.',
    'Refactor a bad design into a Singleton implementation.',
    'Apply Clean Code practices.',
    'Write test cases to verify Singleton behavior.',
]:
    doc.add_paragraph(o, style='List Bullet')

doc.add_heading('Scenario', level=2)
add(doc, 'A college system needs a single configuration class that stores college name and system version.')
add(doc, 'College Name = pcps college')
add(doc, 'System Version = 1.0')

doc.add_heading('Part A: Instructions', level=2)
for i in ['Bad design code', 'Problems in bad design', 'Refactored good design using Singleton Pattern',
           'Clean code explanation', 'Test cases', 'Short conclusion: Where this pattern is useful in real projects']:
    doc.add_paragraph(i, style='List Bullet')

doc.add_heading('Bad Design Code', level=3)
add(doc, 'File: CollegeConfigBad.java')
add_code(doc, read_src('singleton/CollegeConfigBad.java'))
add(doc, 'File: SingletonBadDemo.java')
add_code(doc, read_src('singleton/SingletonBadDemo.java'))

doc.add_heading('Part B: Analysis Questions', level=2)
doc.add_heading('Q1. How many objects are created?', level=3)
add(doc, 'Two objects are created - one for config1 and another for config2. The new keyword is called twice, so the JVM allocates separate memory for each.')
doc.add_heading('Q2. Why does the comparison return false?', level=3)
add(doc, 'System.out.println(config1 == config2); returns false because == checks reference equality. config1 and config2 point to two distinct memory locations since new CollegeConfigBad() was called twice.')
doc.add_heading('Q3. What problems can occur if configuration objects are created repeatedly?', level=3)
for p in ['Wasted memory - each instance duplicates the same configuration data.',
           'Inconsistent state - different parts of the application may hold different configuration values if mutation is allowed.',
           'Performance overhead - repeated object creation increases GC pressure and slows down the system.']:
    doc.add_paragraph(p, style='List Bullet')
doc.add_heading('Q4. Is creating multiple configuration objects necessary in this scenario? Why?', level=3)
add(doc, 'No. The configuration (college name and version) is read-only, global data that never changes during application runtime. A single shared instance is sufficient and efficient.')
doc.add_heading('Q5. Which software quality attributes are affected?', level=3)
add(doc, 'Memory Efficiency (Performance), Consistency (Reliability), and Maintainability (Single Point of Truth).')

doc.add_heading('Part C: Problems in the Existing Design', level=2)
add(doc, '1. Multiple Instances: The public constructor allows unlimited instantiation, breaking the requirement for a single configuration object.')
add(doc, '2. No Global Access Point: There is no static getInstance() method; callers must manage references manually or create new instances.')
add(doc, '3. Violates Single Responsibility: The class only stores data but does not enforce its own lifecycle rule (that only one instance should exist).')

doc.add_heading('Refactored Good Design using Singleton Pattern', level=3)
add(doc, 'Key improvements:')
for i in ['Private constructor prevents external instantiation.',
           'Eagerly created static final instance ensures exactly one object.',
           'Global access point via getInstance() method.']:
    doc.add_paragraph(i, style='List Bullet')
add(doc, 'File: CollegeConfig.java')
add_code(doc, read_src('singleton/CollegeConfig.java'))
add(doc, 'File: SingletonDemo.java')
add_code(doc, read_src('singleton/SingletonDemo.java'))

doc.add_heading('Clean Code Explanation', level=2)
for c in ['Eager Initialization: The instance is created at class load time by the JVM, which is thread-safe without synchronization.',
           'Private Constructor: The constructor is hidden, so no external code can bypass getInstance().',
           'Descriptive Naming: getInstance() clearly conveys the intent - return the single instance.',
           'Minimal Surface Area: The class does only one thing: provide one global configuration instance.',
           'Immutability-Friendly: Fields can be made final if configuration is truly read-once.']:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('Test Cases', level=2)
add(doc, 'File: CollegeConfigTest.java')
add_code(doc, read_src('singleton/CollegeConfigTest.java'))
add(doc, 'Test Descriptions:')
for t in ['testSingleInstance(): Verifies that two calls to getInstance() return the exact same object (reference equality).',
           'testCollegeName(): Verifies the default college name is "pcps college".',
           'testSystemVersion(): Verifies the default system version is "1.0".',
           'testInstanceNotNull(): Ensures getInstance() never returns null.']:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('Conclusion', level=2)
add(doc, 'The Singleton pattern is useful in real projects for managing shared resources such as database connections, thread pools, caches, logging services, and application configuration objects. It guarantees a single point of control and eliminates unnecessary duplication. Common frameworks like Spring use singleton-scoped beans by default.')
doc.add_page_break()

# ===== PRACTICAL 2: FACTORY =====
doc.add_heading('Practical 2: Factory Method Pattern', level=1)
doc.add_heading('Objective', level=2)
for o in [
    'Identify problems caused by creating objects using multiple if-else statements.',
    'Understand the need for the Factory Method Design Pattern.',
    'Refactor a bad design into a Factory Method implementation.',
    'Apply Clean Code practices.',
    'Write test cases to verify Factory Method behavior.',
]:
    doc.add_paragraph(o, style='List Bullet')

doc.add_heading('Scenario', level=2)
add(doc, 'A notification system sends messages using different channels: EMAIL, SMS, PUSH.')
add(doc, 'Instead of creating objects manually using many if-else statements everywhere, object creation should be centralized.')

doc.add_heading('Part A: Instructions', level=2)
for i in ['Bad design code', 'Problems in bad design', 'Refactored good design using Factory Method Pattern',
           'Clean code explanation', 'Test cases', 'Short conclusion: Where this pattern is useful in real projects']:
    doc.add_paragraph(i, style='List Bullet')

doc.add_heading('Bad Design Code', level=3)
add(doc, 'File: NotificationServiceBad.java')
add_code(doc, read_src('factory/NotificationServiceBad.java'))
add(doc, 'File: FactoryBadDemo.java')
add_code(doc, read_src('factory/FactoryBadDemo.java'))

doc.add_heading('Part B: Analysis Questions', level=2)
doc.add_heading('Q1. How does the program decide which notification type to send?', level=3)
add(doc, 'The program uses a chain of if-else statements that check the type parameter. Each branch executes the corresponding notification logic inline.')
doc.add_heading('Q2. Why are multiple if-else statements considered a bad design?', level=3)
for r in ['Violates the Open/Closed Principle - adding a new notification type requires editing the class.',
           'Duplicates logic - if the same type-checking appears in multiple places, maintenance becomes error-prone.',
           'Reduces readability - long if-else chains are hard to scan and understand.']:
    doc.add_paragraph(r, style='List Bullet')
doc.add_heading('Q3. What problems can occur if notification creation logic is written repeatedly?', level=3)
add(doc, 'Code duplication leads to inconsistencies (one branch is updated but another is forgotten), increased testing effort, and higher risk of bugs when adding new notification types.')
doc.add_heading('Q4. Is centralizing object creation necessary in this scenario? Why?', level=3)
add(doc, 'Yes. Centralizing creation in a dedicated factory class means adding a new notification type requires only a new concrete class and a single update to the factory method - no changes to client code.')
doc.add_heading('Q5. Which software quality attributes are affected?', level=3)
add(doc, 'Maintainability (Open/Closed Principle), Extensibility (easy to add new types), and Readability (single creation point).')

doc.add_heading('Part C: Problems in the Existing Design', level=2)
add(doc, '1. Open/Closed Violation: Adding a new type (e.g., WHATSAPP) requires modifying NotificationServiceBad.')
add(doc, '2. Logic Mixed with Creation: Notification logic is tangled with type-dispatching, making the class hard to test and reuse.')
add(doc, '3. No Abstraction: Client code depends on concrete implementation details rather than a common interface.')

doc.add_heading('Refactored Good Design using Factory Pattern', level=3)
add(doc, 'Key improvements:')
for i in ['Notification interface defines the contract - Client depends on abstraction, not concrete classes.',
           'NotificationFactory centralizes object creation - Adding a new type only requires a new class and a factory branch.',
           'Each concrete notification class (EmailNotification, SmsNotification, PushNotification) is a separate, testable unit.']:
    doc.add_paragraph(i, style='List Bullet')
add(doc, 'File: Notification.java')
add_code(doc, read_src('factory/Notification.java'))
add(doc, 'File: EmailNotification.java')
add_code(doc, read_src('factory/EmailNotification.java'))
add(doc, 'File: SmsNotification.java')
add_code(doc, read_src('factory/SmsNotification.java'))
add(doc, 'File: PushNotification.java')
add_code(doc, read_src('factory/PushNotification.java'))
add(doc, 'File: NotificationFactory.java')
add_code(doc, read_src('factory/NotificationFactory.java'))
add(doc, 'File: FactoryDemo.java')
add_code(doc, read_src('factory/FactoryDemo.java'))

doc.add_heading('Clean Code Explanation', level=2)
for c in ['Single Responsibility: NotificationFactory only creates objects. Each Notification subclass only sends messages.',
           'Open/Closed Principle: The factory can be extended with new types without modifying existing notification classes.',
           'Interface-Based Design: Client code depends on the Notification interface, not on concrete implementations (Dependency Inversion).',
           'Descriptive Naming: createNotification() clearly communicates its purpose.']:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('Test Cases', level=2)
add(doc, 'File: NotificationFactoryTest.java')
add_code(doc, read_src('factory/NotificationFactoryTest.java'))
add(doc, 'Test Descriptions:')
for t in ['testEmailNotification(): Verifies factory returns an EmailNotification for "EMAIL".',
           'testSmsNotification(): Verifies factory returns an SmsNotification for "SMS".',
           'testPushNotification(): Verifies factory returns a PushNotification for "PUSH".',
           'testNotificationNotNull(): Ensures the returned object is not null.',
           'testInvalidTypeThrowsException(): Verifies IllegalArgumentException for unknown types.']:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('Conclusion', level=2)
add(doc, 'The Factory Method Pattern is widely used in real projects for managing object creation in logging frameworks (SLF4J), GUI toolkits (creating platform-specific widgets), payment gateway integrations (different processors), and data access layers (different database drivers). It decouples client code from concrete classes, making systems more extensible and maintainable.')
doc.add_page_break()

# ===== PRACTICAL 3: BUILDER =====
doc.add_heading('Practical 3: Builder Pattern', level=1)
doc.add_heading('Objective', level=2)
for o in [
    'Identify problems caused by constructors with many parameters.',
    'Understand the need for the Builder Design Pattern.',
    'Refactor a bad design into a Builder Pattern implementation.',
    'Apply Clean Code practices.',
    'Write test cases to verify Builder Pattern behavior.',
]:
    doc.add_paragraph(o, style='List Bullet')

doc.add_heading('Scenario', level=2)
add(doc, 'A student registration module creates student objects with the following fields: studentId, name, email, phoneNumber, department, semester, address, guardianName.')
add(doc, 'Creating such an object using a long constructor becomes confusing and error-prone.')

doc.add_heading('Part A: Instructions', level=2)
for i in ['Bad design code', 'Problems in bad design', 'Refactored good design using Builder Pattern',
           'Clean code explanation', 'Test cases', 'Short conclusion: Where this pattern is useful in real projects']:
    doc.add_paragraph(i, style='List Bullet')

doc.add_heading('Bad Design Code', level=3)
add(doc, 'File: StudentBad.java')
add_code(doc, read_src('builder/StudentBad.java'))
add(doc, 'File: BuilderBadDemo.java')
add_code(doc, read_src('builder/BuilderBadDemo.java'))

doc.add_heading('Part B: Analysis Questions', level=2)
doc.add_heading('Q1. How is the Student object created in the current implementation?', level=3)
add(doc, 'It is created by passing all 8 arguments to a single constructor. The caller must remember the exact order and types of every parameter.')
doc.add_heading('Q2. Why is using a constructor with many parameters considered a bad design?', level=3)
for r in ['Reduces readability - it is impossible to tell which argument maps to which field without counting positions.',
           'Increases error-proneness - swapping two String arguments (e.g., email and phone) compiles but produces incorrect state.',
           'Forces all fields at once - there is no way to create an object with only some fields set without adding multiple constructors (telescoping anti-pattern).']:
    doc.add_paragraph(r, style='List Bullet')
doc.add_heading('Q3. What problems can occur if constructors have many parameters?', level=3)
add(doc, 'Developers may inadvertently swap arguments, forget fields, or create telescoping constructors. The code becomes brittle; adding or reordering parameters breaks all call sites.')
doc.add_heading('Q4. Is using the Builder Pattern necessary in this scenario? Why?', level=3)
add(doc, 'Yes. With 8 fields (2 int + 6 String), the Builder Pattern makes object creation self-documenting through named setter methods. It also supports partial construction and optional fields naturally.')
doc.add_heading('Q5. Which software quality attributes are affected?', level=3)
add(doc, 'Readability (named parameters), Maintainability (adding fields does not break existing code), and Reliability (fewer argument-swap bugs).')

doc.add_heading('Part C: Problems in the Existing Design', level=2)
add(doc, '1. Positional Confusion: The caller must remember the exact order of 8 parameters - swapping email and phoneNumber compiles but is logically wrong.')
add(doc, '2. No Partial Construction: The only way to create a StudentBad is with all 8 fields; optional fields force null or dummy values.')
add(doc, '3. Poor Readability: new StudentBad(101, "Ram", "ram@gmail.com", ...) conveys no semantic meaning about which argument is which.')

doc.add_heading('Refactored Good Design using Builder Pattern', level=3)
add(doc, 'Key improvements:')
for i in ['Fluent method chaining with named setters makes each field explicit.',
           'Inner static Builder class isolates construction logic.',
           'Immutable Student object - all fields are final, built once via private constructor.']:
    doc.add_paragraph(i, style='List Bullet')
add(doc, 'File: Student.java')
add_code(doc, read_src('builder/Student.java'))
add(doc, 'File: BuilderDemo.java')
add_code(doc, read_src('builder/BuilderDemo.java'))

doc.add_heading('Clean Code Explanation', level=2)
for c in ['Fluent API: Each setter returns this, enabling readable method chaining.',
           'Immutability: All Student fields are final, making the object thread-safe and predictable.',
           'Separation of Concerns: The Builder handles incremental construction; Student only exposes data.',
           'Named Parameters: .name("Ram") is self-documenting compared to positional "Ram" in a constructor.',
           'Partial Construction: The Builder allows setting only the needed fields without telescoping constructors.']:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('Test Cases', level=2)
add(doc, 'File: StudentBuilderTest.java')
add_code(doc, read_src('builder/StudentBuilderTest.java'))
add(doc, 'Test Descriptions:')
for t in ['testStudentCreatedSuccessfully(): Verifies a fully built Student object is not null.',
           'testStudentId(): Verifies studentId is set correctly via Builder.',
           'testStudentName(): Verifies name is set correctly.',
           'testDepartment(): Verifies department is set correctly.',
           'testSemester(): Verifies semester is set correctly.',
           'testEmail(): Verifies email is set correctly.',
           'testPhoneNumber(): Verifies phoneNumber is set correctly.',
           'testAddress(): Verifies address is set correctly.',
           'testGuardianName(): Verifies guardianName is set correctly.']:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('Conclusion', level=2)
add(doc, 'The Builder Pattern is widely used in real projects for constructing complex objects in Java, such as StringBuilder/StringBuffer, Spring RestTemplate builders, Lombok @Builder, and Google Protocol Buffers message builders. It is essential when an object has many optional fields, when construction requires multiple steps, or when immutability is desired.')
doc.add_page_break()

# ===== PRACTICAL 4: ADAPTER =====
doc.add_heading('Practical 4: Adapter Pattern', level=1)
doc.add_heading('Objective', level=2)
for o in [
    'Identify problems caused by incompatible interfaces.',
    'Understand the need for the Adapter Design Pattern.',
    'Refactor a bad design into an Adapter implementation.',
    'Apply Clean Code practices.',
    'Write test cases to verify Adapter behavior.',
]:
    doc.add_paragraph(o, style='List Bullet')

doc.add_heading('Scenario', level=2)
add(doc, 'PCPS College has developed a new Student Notification System.')
add(doc, 'The system expects all notification services to implement: sendNotification(String message);')
add(doc, 'However, an old email service already exists: sendMail(String text);')
add(doc, 'The new system must use the old service without modifying its code.')

doc.add_heading('Part A: Instructions', level=2)
for i in ['Bad design code', 'Problems in bad design', 'Refactored good design using Adapter Pattern',
           'Clean code explanation', 'Test cases', 'Short conclusion: Where this pattern is useful in real projects']:
    doc.add_paragraph(i, style='List Bullet')

doc.add_heading('Bad Design Code', level=3)
add(doc, 'File: AdapterBadDemo.java')
add_code(doc, read_src('adapter/AdapterBadDemo.java'))

doc.add_heading('Part B: Analysis Questions', level=2)
doc.add_heading('Q1. Why cant the notification system directly use the legacy email service?', level=3)
add(doc, 'Because the legacy service uses sendMail(String text) while the new system expects sendNotification(String message). The interfaces are incompatible.')
doc.add_heading('Q2. What problems arise when two classes have different interfaces?', level=3)
for r in ['Code cannot be reused without modification.',
           'Developers must write custom glue code in every client that needs the incompatible service.',
           'Testing becomes harder because there is no standard contract between components.']:
    doc.add_paragraph(r, style='List Bullet')
doc.add_heading('Q3. What happens if multiple old systems need integration?', level=3)
add(doc, 'Each integration requires its own adapter, but the overall architecture remains clean because the adapters are isolated, reusable classes that follow a common interface.')
doc.add_heading('Q4. Is modifying third-party code a good practice? Why?', level=3)
add(doc, 'No. Third-party code should not be modified because changes are lost on updates, the original code may be closed-source, and modifying it introduces maintenance coupling. The Adapter pattern avoids this entirely.')
doc.add_heading('Q5. Which software quality attributes are affected?', level=3)
add(doc, 'Reusability (adapt old code without changes), Maintainability (changes isolated to the adapter), and Flexibility (swap implementations easily).')

doc.add_heading('Part C: Problems in the Existing Design', level=2)
add(doc, '1. Interface Mismatch: The new system requires sendNotification(), but the legacy service only provides sendMail(). Direct use is impossible without changes.')
add(doc, '2. No Code Reuse: Without an adapter, developers cannot reuse the existing LegacyEmailService in the new system.')
add(doc, '3. Violates Open/Closed: If third-party code must be modified to fit the new interface, it violates the Open/Closed Principle.')

doc.add_heading('Refactored Good Design using Adapter Pattern', level=3)
add(doc, 'Key improvements:')
for i in ['EmailAdapter wraps LegacyEmailService and implements the Notification interface.',
           'No changes to LegacyEmailService - third-party code remains untouched.',
           'Client code depends only on the Notification interface, not on LegacyEmailService.']:
    doc.add_paragraph(i, style='List Bullet')
add(doc, 'File: Notification.java')
add_code(doc, read_src('adapter/Notification.java'))
add(doc, 'File: LegacyEmailService.java')
add_code(doc, read_src('adapter/LegacyEmailService.java'))
add(doc, 'File: EmailAdapter.java')
add_code(doc, read_src('adapter/EmailAdapter.java'))
add(doc, 'File: AdapterDemo.java')
add_code(doc, read_src('adapter/AdapterDemo.java'))

doc.add_heading('Clean Code Explanation', level=2)
for c in ['Single Responsibility: EmailAdapter only handles interface conversion.',
           'Open/Closed Principle: The legacy service is not modified; the adapter extends its use.',
           'Interface Segregation: The Notification interface is minimal and focused.',
           'Composition over Inheritance: Adapter wraps the adaptee rather than subclassing it.']:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('Test Cases', level=2)
add(doc, 'File: AdapterTest.java')
add_code(doc, read_src('adapter/AdapterTest.java'))
add(doc, 'Test Descriptions:')
for t in ['testAdapterObjectCreated(): Verifies the adapter object is created successfully.',
           'testAdapterDelegatesCorrectly(): Verifies sendNotification() delegates to sendMail() without exceptions.',
           'testAdapterImplementsNotification(): Verifies the adapter implements the Notification interface.',
           'testAdapterNotNull(): Ensures the adapter object is not null.']:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('Conclusion', level=2)
add(doc, 'The Adapter Pattern is widely used in real projects for integrating legacy systems, wrapping third-party libraries, and making incompatible interfaces work together. Examples include JDBC drivers (adapting database-specific APIs to a common SQL interface), Java Collections (Arrays.asList()), and GUI frameworks (adapting older event models to newer ones).')
doc.add_page_break()

# ===== PRACTICAL 5: FACADE =====
doc.add_heading('Practical 5: Facade Pattern', level=1)
doc.add_heading('Objective', level=2)
for o in [
    'Identify problems caused by complex subsystem interactions.',
    'Understand the need for the Facade Design Pattern.',
    'Refactor a bad design into a Facade implementation.',
    'Apply Clean Code practices.',
    'Write test cases to verify Facade behavior.',
]:
    doc.add_paragraph(o, style='List Bullet')

doc.add_heading('Scenario', level=2)
add(doc, 'Student registration requires: Validation, Roll Number Generation, Save Student, Department Assignment, Notification.')
add(doc, 'Currently, every service is called manually.')

doc.add_heading('Part A: Instructions', level=2)
for i in ['Bad design code', 'Problems in bad design', 'Refactored good design using Facade Pattern',
           'Clean code explanation', 'Test cases', 'Short conclusion']:
    doc.add_paragraph(i, style='List Bullet')

doc.add_heading('Bad Design Code', level=3)
add(doc, 'File: FacadeBadDemo.java')
add_code(doc, read_src('facade/FacadeBadDemo.java'))

doc.add_heading('Part B: Analysis Questions', level=2)
doc.add_heading('Q1. How many services are required to register a student?', level=3)
add(doc, 'Five services are required: StudentValidator, RollNumberAssigner, DepartmentAssigner, StudentSaver, and WelcomeMessageSender.')
doc.add_heading('Q2. Why is the client tightly coupled with multiple services?', level=3)
add(doc, 'The client must instantiate each service, know the correct call order, and handle failures for each step. Any change to the workflow requires updating every client.')
doc.add_heading('Q3. What happens if registration process changes?', level=3)
add(doc, 'Every client that manually orchestrates the services must be updated. This is error-prone and increases maintenance cost.')
doc.add_heading('Q4. How does Facade simplify the client code?', level=3)
add(doc, 'The facade provides a single registerStudent() method that internally calls all five services. The client only depends on one class instead of five.')
doc.add_heading('Q5. Which software quality attributes are affected?', level=3)
add(doc, 'Maintainability (single change point), Loose Coupling (client depends only on facade), and Simplicity (reduced client complexity).')

doc.add_heading('Part C: Problems in the Existing Design', level=2)
add(doc, '1. Tight Coupling: The client is tightly coupled to all five subsystem classes.')
add(doc, '2. Scattered Workflow Logic: The registration sequence is duplicated wherever a student is registered.')
add(doc, '3. High Maintenance Cost: Adding or reordering a registration step requires editing every client.')

doc.add_heading('Refactored Good Design using Facade Pattern', level=3)
add(doc, 'Key improvements:')
for i in ['StudentRegistrationFacade provides a single registerStudent() method.',
           'Client code is decoupled from the individual subsystem classes.',
           'The registration workflow is centralized and easy to maintain.']:
    doc.add_paragraph(i, style='List Bullet')
add(doc, 'File: StudentValidator.java')
add_code(doc, read_src('facade/StudentValidator.java'))
add(doc, 'File: RollNumberAssigner.java')
add_code(doc, read_src('facade/RollNumberAssigner.java'))
add(doc, 'File: DepartmentAssigner.java')
add_code(doc, read_src('facade/DepartmentAssigner.java'))
add(doc, 'File: StudentSaver.java')
add_code(doc, read_src('facade/StudentSaver.java'))
add(doc, 'File: WelcomeMessageSender.java')
add_code(doc, read_src('facade/WelcomeMessageSender.java'))
add(doc, 'File: StudentRegistrationFacade.java')
add_code(doc, read_src('facade/StudentRegistrationFacade.java'))
add(doc, 'File: FacadeDemo.java')
add_code(doc, read_src('facade/FacadeDemo.java'))

doc.add_heading('Clean Code Explanation', level=2)
for c in ['Single Responsibility: The facade coordinates subsystems; each subsystem has one job.',
           'Minimal Client Dependency: Client code only depends on the facade class.',
           'Simplified API: A complex multi-step workflow is exposed as a single method call.',
           'Easy to Test: The facade can be tested independently; subsystems can be tested separately.']:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('Test Cases', level=2)
add(doc, 'File: FacadeTest.java')
add_code(doc, read_src('facade/FacadeTest.java'))
add(doc, 'Test Descriptions:')
for t in ['testRegistrationStartsSuccessfully(): Verifies registration starts without exceptions.',
           'testValidationServiceExecuted(): Verifies validation service is executed (throws for invalid data).',
           'testDepartmentAllocationExecuted(): Verifies department allocation runs successfully.',
           'testNotificationServiceExecuted(): Verifies notification service runs successfully.',
           'testFacadeNotNull(): Ensures the facade object is not null.']:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('Conclusion', level=2)
add(doc, 'The Facade Pattern is widely used in real projects to simplify complex APIs. Examples include Spring JdbcTemplate (simplifies JDBC), SLF4J (unified logging facade), and media players that hide decoding, rendering, and audio subsystems.')
doc.add_page_break()

# ===== PRACTICAL 6: PROXY =====
doc.add_heading('Practical 6: Proxy Pattern', level=1)
doc.add_heading('Objective', level=2)
for o in [
    'Identify problems caused by unrestricted access.',
    'Understand the need for the Proxy Pattern.',
    'Refactor a bad design into a Proxy implementation.',
    'Apply Clean Code practices.',
    'Write test cases to verify Proxy behavior.',
]:
    doc.add_paragraph(o, style='List Bullet')

doc.add_heading('Scenario', level=2)
add(doc, 'Only Admin can view student marks.')
add(doc, 'Role = ADMIN = Allowed.')
add(doc, 'Role = STUDENT = Denied.')

doc.add_heading('Part A: Instructions', level=2)
for i in ['Bad design code', 'Problems in bad design', 'Refactored good design using Proxy Pattern',
           'Clean code explanation', 'Test cases', 'Short conclusion']:
    doc.add_paragraph(i, style='List Bullet')

doc.add_heading('Bad Design Code', level=3)
add(doc, 'File: ProxyBadDemo.java')
add_code(doc, read_src('proxy/ProxyBadDemo.java'))

doc.add_heading('Part B: Analysis Questions', level=2)
doc.add_heading('Q1. Who can currently access student records?', level=3)
add(doc, 'Currently, anyone with access to RealStudentRecord can view marks. Access control logic is scattered in client code rather than enforced by the class itself.')
doc.add_heading('Q2. Why is this a security issue?', level=3)
for r in ['Access checks can be forgotten or bypassed by new clients.',
           'The RealStudentRecord has no protection mechanism of its own.',
           'Sensitive data is exposed to unauthorized roles.']:
    doc.add_paragraph(r, style='List Bullet')
doc.add_heading('Q3. What happens if access checks are repeated everywhere?', level=3)
add(doc, 'Code duplication leads to inconsistencies - one developer may forget the check, use a different role string, or implement weaker validation.')
doc.add_heading('Q4. Why should security be centralized?', level=3)
add(doc, 'Centralized security in the proxy ensures consistent enforcement, easier auditing, and a single point of change when security policies are updated.')
doc.add_heading('Q5. Which software quality attributes are affected?', level=3)
add(doc, 'Security (centralized access control), Maintainability (single point of change), and Consistency (uniform enforcement across all clients).')

doc.add_heading('Part C: Problems in the Existing Design', level=2)
add(doc, '1. No Access Control: RealStudentRecord has no built-in protection - anyone can call viewMarks().')
add(doc, '2. Scattered Security Logic: Every client must duplicate the role-checking code.')
add(doc, '3. Tight Coupling: Client code is coupled to both RealStudentRecord and the access control logic.')

doc.add_heading('Refactored Good Design using Proxy Pattern', level=3)
add(doc, 'Key improvements:')
for i in ['StudentRecordProxy controls access to RealStudentRecord.',
           'Access control logic is centralized in the proxy, not scattered in client code.',
           'Client code depends only on the StudentRecord interface.']:
    doc.add_paragraph(i, style='List Bullet')
add(doc, 'File: StudentRecord.java')
add_code(doc, read_src('proxy/StudentRecord.java'))
add(doc, 'File: RealStudentRecord.java')
add_code(doc, read_src('proxy/RealStudentRecord.java'))
add(doc, 'File: StudentRecordProxy.java')
add_code(doc, read_src('proxy/StudentRecordProxy.java'))
add(doc, 'File: ProxyDemo.java')
add_code(doc, read_src('proxy/ProxyDemo.java'))

doc.add_heading('Clean Code Explanation', level=2)
for c in ['Single Responsibility: Proxy handles access control; RealStudentRecord handles data.',
           'Open/Closed Principle: New proxy behavior (caching, logging) can be added without changing the real subject.',
           'Interface-Based Design: Client depends on StudentRecord interface, not on the proxy or real subject.',
           'Lazy Initialization: The proxy can defer creating the real subject until absolutely needed.']:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('Test Cases', level=2)
add(doc, 'File: ProxyTest.java')
add_code(doc, read_src('proxy/ProxyTest.java'))
add(doc, 'Test Descriptions:')
for t in ['testAdminAccessAllowed(): Verifies ADMIN role is allowed to view marks.',
           'testStudentAccessDenied(): Verifies STUDENT role is denied.',
           'testProxyObjectCreated(): Verifies the proxy implements the StudentRecord interface.',
           'testProxyNotNull(): Ensures the proxy object is not null.']:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('Conclusion', level=2)
add(doc, 'The Proxy Pattern is widely used in real projects for access control (security proxies), lazy loading (virtual proxies), caching (cache proxies), logging, and remote method invocation (RMI proxies). Examples include Spring AOP proxies, Hibernate lazy-loading proxies, and Java RMI stubs.')
doc.add_page_break()

# ===== PRACTICAL 7: DECORATOR =====
doc.add_heading('Practical 7: Decorator Pattern', level=1)
doc.add_heading('Objective', level=2)
for o in [
    'Identify problems caused by excessive inheritance.',
    'Understand the need for the Decorator Pattern.',
    'Refactor a bad design into a Decorator implementation.',
    'Apply Clean Code practices.',
    'Write test cases to verify Decorator behavior.',
]:
    doc.add_paragraph(o, style='List Bullet')

doc.add_heading('Scenario', level=2)
add(doc, 'The Student Report System generates student reports.')
add(doc, 'A basic report contains student ID and name.')
add(doc, 'Additional features (PDF export, Watermark, Digital Signature) can be added dynamically.')

doc.add_heading('Part A: Instructions', level=2)
for i in ['Bad design code', 'Problems in bad design', 'Refactored good design using Decorator Pattern',
           'Clean code explanation', 'Test cases', 'Short conclusion']:
    doc.add_paragraph(i, style='List Bullet')

doc.add_heading('Bad Design Code', level=3)
add(doc, 'File: DecoratorBadDemo.java')
add_code(doc, read_src('decorator/DecoratorBadDemo.java'))

doc.add_heading('Part B: Analysis Questions', level=2)
doc.add_heading('Q1. Why are many report classes being created?', level=3)
add(doc, 'Without the Decorator pattern, every combination of features (PDF, Watermark, Signature) requires a separate subclass, leading to class explosion (2^N classes for N features).')
doc.add_heading('Q2. What happens when a new feature is added?', level=3)
add(doc, 'With inheritance, adding one new feature doubles the number of subclasses (every existing combination needs a new version with the feature).')
doc.add_heading('Q3. Why does inheritance become difficult to maintain?', level=3)
add(doc, 'Deep inheritance hierarchies are fragile: a change in a base class can break many subclasses, and understanding the complete behavior requires tracing through multiple levels.')
doc.add_heading('Q4. How can Decorator add features dynamically?', level=3)
add(doc, 'The Decorator pattern uses composition: each decorator wraps a StudentReport and adds its behavior before/after delegating to the wrapped object.')
doc.add_heading('Q5. Which software quality attributes are affected?', level=3)
add(doc, 'Flexibility (runtime feature composition), Maintainability (no modification of existing code), and Extensibility (new decorators added independently).')

doc.add_heading('Part C: Problems in the Existing Design', level=2)
add(doc, '1. Class Explosion: With N features, inheritance requires up to 2^N subclasses.')
add(doc, '2. Rigid at Compile Time: Feature combinations are fixed at compile time.')
add(doc, '3. Duplicated Code: Common behavior must be repeated across multiple subclasses.')

doc.add_heading('Refactored Good Design using Decorator Pattern', level=3)
add(doc, 'Key improvements:')
for i in ['BasicStudentReport remains unchanged - new features are added via decorators.',
           'Decorators can be stacked in any order at runtime.',
           'Each decorator handles one concern (single responsibility).']:
    doc.add_paragraph(i, style='List Bullet')
add(doc, 'File: StudentReport.java')
add_code(doc, read_src('decorator/StudentReport.java'))
add(doc, 'File: BasicStudentReport.java')
add_code(doc, read_src('decorator/BasicStudentReport.java'))
add(doc, 'File: ReportDecorator.java')
add_code(doc, read_src('decorator/ReportDecorator.java'))
add(doc, 'File: PdfDecorator.java')
add_code(doc, read_src('decorator/PdfDecorator.java'))
add(doc, 'File: WatermarkDecorator.java')
add_code(doc, read_src('decorator/WatermarkDecorator.java'))
add(doc, 'File: SignatureDecorator.java')
add_code(doc, read_src('decorator/SignatureDecorator.java'))
add(doc, 'File: DecoratorDemo.java')
add_code(doc, read_src('decorator/DecoratorDemo.java'))

doc.add_heading('Clean Code Explanation', level=2)
for c in ['Single Responsibility: Each decorator adds exactly one feature.',
           'Open/Closed Principle: Core classes are never modified; extension happens via new decorator classes.',
           'Composition over Inheritance: Behaviors are combined by wrapping, not by subclassing.',
           'Runtime Flexibility: The set of active features is determined at object construction time.']:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('Test Cases', level=2)
add(doc, 'File: DecoratorTest.java')
add_code(doc, read_src('decorator/DecoratorTest.java'))
add(doc, 'Test Descriptions:')
for t in ['testBasicReportGeneration(): Verifies basic report generation works.',
           'testPdfFeatureWorks(): Verifies PDF feature works when decorated.',
           'testWatermarkFeatureWorks(): Verifies Watermark feature works when decorated.',
           'testSignatureFeatureWorks(): Verifies Signature feature works when decorated.',
           'testDecoratedReportNotNull(): Ensures the decorated report is not null.']:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('Conclusion', level=2)
add(doc, 'The Decorator Pattern is widely used in real projects for adding behavior dynamically. Examples include Java I/O streams (BufferedInputStream wraps FileInputStream), GUI toolkits (adding scrollbars, borders to components), and web application middleware (adding logging, authentication, compression).')
doc.add_page_break()

# ===== ASSIGNMENT 1 =====
doc.add_heading('Assignment 1: Student Admission Mini Module', level=1)
doc.add_heading('Objective', level=2)
for o in [
    'Integrate multiple Creational Design Patterns into a single application.',
    'Reuse previously implemented design patterns.',
    'Understand how Singleton, Builder, and Factory Method collaborate.',
    'Apply Clean Code practices.',
    'Develop a simple Student Admission workflow.',
]:
    doc.add_paragraph(o, style='List Bullet')

doc.add_heading('Background', level=2)
add(doc, 'The Admission Module supports:')
for s in ['Loading college configuration information (Singleton).',
           'Creating student profiles (Builder).',
           'Sending welcome notifications (Factory).']:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('Important Instructions', level=2)
for imp in ['Reuse existing implementations from Practical 1, 2, 3.',
            'Do NOT rewrite the design pattern implementations.',
            'Do NOT create new versions of the same patterns.',
            'Only create the integration classes required.']:
    doc.add_paragraph(imp, style='List Number')

doc.add_heading('Integration Code', level=2)
add(doc, 'File: AdmissionService.java')
add_code(doc, read_src('admission/AdmissionService.java'))
add(doc, 'File: AdmissionDemo.java')
add_code(doc, read_src('admission/AdmissionDemo.java'))
add(doc, 'File: AdmissionTest.java')
add_code(doc, read_src('admission/AdmissionTest.java'))

doc.add_heading('How the Patterns Collaborate', level=2)
for c in ['Singleton (CollegeConfig): Single configuration instance for the admission process.',
           'Builder (Student.Builder): Constructs complex Student objects with fluent chaining.',
           'Factory (NotificationFactory): Creates notifications without coupling to concrete classes.']:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('Conclusion', level=2)
add(doc, 'This module demonstrates three creational patterns working together in an admission workflow. The Singleton ensures consistent configuration, the Builder creates profiles cleanly, and the Factory decouples notification creation.')
doc.add_page_break()

# ===== ASSIGNMENT 2 =====
doc.add_heading('Assignment 2: Student Services Mini Module', level=1)
doc.add_heading('Objective', level=2)
for o in [
    'Integrate multiple Structural Design Patterns into a single application.',
    'Reuse previously implemented design patterns.',
    'Understand how Adapter, Facade, Proxy, and Decorator collaborate.',
    'Apply Clean Code practices.',
    'Extend the Admission System into a Student Services System.',
]:
    doc.add_paragraph(o, style='List Bullet')

doc.add_heading('Background', level=2)
add(doc, 'In the previous assignment, students developed a Student Admission Mini Module using Creational Design Patterns (Singleton, Factory Method, Builder).')
add(doc, 'In this assignment, students will extend the system into a Student Services Module by integrating Structural Design Patterns.')

doc.add_heading('Scenario', level=2)
add(doc, 'PCPS College is developing a Student Services Module.')
add(doc, 'When a student is registered:')
for s in ['Registration process is simplified using a Facade.',
           'Welcome email is sent using a legacy email system through an Adapter.',
           'Student records are protected using a Proxy.',
           'Student reports can be enhanced using Decorators.']:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('Important Instructions', level=2)
for imp in ['Reuse implementations from Practical 4 (Adapter), 5 (Facade), 6 (Proxy), 7 (Decorator).',
            'Do NOT rewrite the design pattern implementations.',
            'Only create the integration classes required.']:
    doc.add_paragraph(imp, style='List Number')

doc.add_heading('Integration Code', level=2)
add(doc, 'File: StudentServicesFacade.java')
add_code(doc, read_src('services/StudentServicesFacade.java'))
add(doc, 'File: ServicesDemo.java')
add_code(doc, read_src('services/ServicesDemo.java'))
add(doc, 'File: ServicesTest.java')
add_code(doc, read_src('services/ServicesTest.java'))

doc.add_heading('How the Patterns Collaborate', level=2)
for c in ['Facade (StudentRegistrationFacade): Simplifies multi-step registration into a single method.',
           'Adapter (EmailAdapter): Bridges incompatible LegacyEmailService to the Notification interface.',
           'Proxy (StudentRecordProxy): Protects sensitive marks data by allowing only ADMIN access.',
           'Decorator (PdfDecorator, WatermarkDecorator, SignatureDecorator): Dynamically adds PDF, watermark, and signature to reports.']:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('Conclusion', level=2)
add(doc, 'This module demonstrates four structural patterns working together. The Facade simplifies registration, the Adapter enables legacy integration, the Proxy secures records, and the Decorator enriches reports. Each pattern addresses a specific structural concern, resulting in a flexible, secure, and maintainable system.')

# ===== SAVE =====
out = r'C:\Users\DeathAngel\Documents\Java\Creational Design Pattern\Creational_Design_Patterns_Report.docx'
doc.save(out)
print(f'Document saved to: {out}')
